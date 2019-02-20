import docx

def lector(dir):
    doc=docx.Document(dir)
    n=len(doc.tables)
    for n_t in n:
        table=doc.tables[n_t]
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph.text)
